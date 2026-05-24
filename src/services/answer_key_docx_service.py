import asyncio
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from src.schemas import AnswerKeySchema, AnswerKeySectionSchema

_HTML_TAG_RE: re.Pattern[str] = re.compile(r"<[^>]+>")

_TITLE_COLOR: RGBColor = RGBColor(0x1E, 0x3A, 0x8A)
_SECTION_COLOR: RGBColor = RGBColor(0x25, 0x63, 0xEB)
_QUESTION_COLOR: RGBColor = RGBColor(0x0F, 0x17, 0x2A)
_ANSWER_COLOR: RGBColor = RGBColor(0x33, 0x41, 0x55)
_MUTED_COLOR: RGBColor = RGBColor(0x6B, 0x72, 0x80)


class AnswerKeyDocxService:
	"""Render an `AnswerKeySchema` as a Word document."""

	async def render(self, answer_key: AnswerKeySchema) -> bytes:
		"""Render the answer key as a DOCX file.

		Args:
		    answer_key: The parsed answer key.

		Returns:
		    The DOCX file as bytes.
		"""
		return await asyncio.to_thread(self._render_sync, answer_key)

	def _render_sync(self, answer_key: AnswerKeySchema) -> bytes:
		"""Synchronously build the DOCX file in memory.

		Args:
		    answer_key: The parsed answer key.

		Returns:
		    The DOCX bytes.
		"""
		document: Document = Document()

		title_paragraph = document.add_paragraph()
		title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
		title_run = title_paragraph.add_run(answer_key.title)
		title_run.font.size = Pt(24)
		title_run.font.bold = True
		title_run.font.color.rgb = _TITLE_COLOR

		if answer_key.subtitle:
			subtitle_paragraph = document.add_paragraph()
			subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
			subtitle_run = subtitle_paragraph.add_run(answer_key.subtitle)
			subtitle_run.font.size = Pt(11)
			subtitle_run.font.color.rgb = _MUTED_COLOR
			subtitle_run.italic = True

		for section in answer_key.sections:
			self._add_section(document, section)

		buffer: io.BytesIO = io.BytesIO()
		document.save(buffer)
		return buffer.getvalue()

	def _add_section(
		self,
		document: Document,
		section: AnswerKeySectionSchema,
	) -> None:
		"""Append a section heading and its items to the document.

		Args:
		    document: The DOCX document under construction.
		    section: The section to render.
		"""
		document.add_paragraph()
		heading = document.add_paragraph()
		heading_run = heading.add_run(section.title)
		heading_run.font.size = Pt(14)
		heading_run.font.bold = True
		heading_run.font.color.rgb = _SECTION_COLOR

		for item in section.items:
			question_paragraph = document.add_paragraph()
			question_run = question_paragraph.add_run(
				_strip_inline_html(item.question),
			)
			question_run.font.size = Pt(12)
			question_run.font.bold = True
			question_run.font.color.rgb = _QUESTION_COLOR
			question_paragraph.paragraph_format.space_before = Pt(8)

			answer_paragraph = document.add_paragraph()
			answer_paragraph.paragraph_format.left_indent = Pt(18)
			label_run = answer_paragraph.add_run("Resposta: ")
			label_run.font.size = Pt(12)
			label_run.font.bold = True
			label_run.font.color.rgb = _SECTION_COLOR
			answer_run = answer_paragraph.add_run(
				_strip_inline_html(item.answer),
			)
			answer_run.font.size = Pt(12)
			answer_run.italic = True
			answer_run.font.color.rgb = _ANSWER_COLOR


def _strip_inline_html(text: str) -> str:
	"""Remove inline HTML tags from a string.

	DOCX cannot render arbitrary HTML, so the simplest accurate path is
	to strip tags and keep only the visible text content.

	Args:
	    text: The source text, possibly containing inline HTML.

	Returns:
	    The text without HTML tags.
	"""
	return _HTML_TAG_RE.sub("", text)
