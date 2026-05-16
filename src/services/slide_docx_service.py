import asyncio
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from src.schemas import SlideDeckSchema, SlideKind, SlideSchema

_TITLE_COLOR: RGBColor = RGBColor(0x1F, 0x29, 0x37)
_ACCENT_COLOR: RGBColor = RGBColor(0x25, 0x63, 0xEB)
_MUTED_COLOR: RGBColor = RGBColor(0x6B, 0x72, 0x80)
_NOTES_COLOR: RGBColor = RGBColor(0x33, 0x41, 0x55)


class SlideDocxService:
    """Render a `SlideDeckSchema` as a Word handout (1 slide per page)."""

    async def render(self, deck: SlideDeckSchema) -> bytes:
        """Render the deck as a DOCX handout.

        Args:
            deck: The parsed slide deck.

        Returns:
            The DOCX file as bytes.
        """
        return await asyncio.to_thread(self._render_sync, deck)

    def _render_sync(self, deck: SlideDeckSchema) -> bytes:
        """Synchronously build the DOCX handout in memory.

        Args:
            deck: The parsed slide deck.

        Returns:
            The DOCX bytes.
        """
        document: Document = Document()

        cover = document.add_paragraph()
        cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_run = cover.add_run(deck.title)
        cover_run.font.size = Pt(28)
        cover_run.font.bold = True
        cover_run.font.color.rgb = _TITLE_COLOR

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Material de apoio do apresentador")
        subtitle_run.font.size = Pt(12)
        subtitle_run.font.color.rgb = _MUTED_COLOR
        subtitle_run.italic = True

        for index, slide in enumerate(deck.slides):
            self._add_page_break(document)
            self._add_slide_page(document, slide, deck)
            if index == len(deck.slides) - 1:
                continue

        buffer: io.BytesIO = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _add_slide_page(
        self,
        document: Document,
        slide: SlideSchema,
        deck: SlideDeckSchema,
    ) -> None:
        """Render one slide as a single handout page.

        Args:
            document: The DOCX document under construction.
            slide: The slide to render.
            deck: The full deck (used for page header information).
        """
        eyebrow_text: str = self._build_eyebrow(slide)
        if eyebrow_text:
            eyebrow = document.add_paragraph()
            eyebrow_run = eyebrow.add_run(eyebrow_text)
            eyebrow_run.font.size = Pt(11)
            eyebrow_run.font.color.rgb = _ACCENT_COLOR
            eyebrow_run.font.bold = True

        title_paragraph = document.add_paragraph()
        title_run = title_paragraph.add_run(slide.title)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = _TITLE_COLOR

        if slide.kind == SlideKind.LESSON_TITLE:
            self._add_subtle_paragraph(
                document,
                "Slide de capa.",
            )
        elif slide.kind == SlideKind.BULLETS:
            for bullet in slide.bullets:
                paragraph = document.add_paragraph(style="List Bullet")
                if bullet.term:
                    term_run = paragraph.add_run(f"{bullet.term}: ")
                    term_run.font.bold = True
                    term_run.font.size = Pt(12)
                description_run = paragraph.add_run(bullet.description)
                description_run.font.size = Pt(12)
        else:
            for paragraph_text in slide.paragraphs:
                paragraph = document.add_paragraph()
                run = paragraph.add_run(paragraph_text)
                run.font.size = Pt(12)
                paragraph.paragraph_format.space_after = Pt(8)

        self._add_notes_section(document, slide)
        self._add_footer_line(document, slide, deck)

    def _add_notes_section(
        self,
        document: Document,
        slide: SlideSchema,
    ) -> None:
        """Append the speaker-notes block to the current page.

        The block always renders so the handout has a place for hand
        annotations even when no notes were provided in the source text.

        Args:
            document: The DOCX document under construction.
            slide: The slide schema providing the optional notes text.
        """
        document.add_paragraph()
        header = document.add_paragraph()
        header_run = header.add_run("Notas do apresentador")
        header_run.font.size = Pt(11)
        header_run.font.bold = True
        header_run.font.color.rgb = _ACCENT_COLOR

        if slide.notes:
            for line in slide.notes.splitlines():
                stripped: str = line.strip()
                if not stripped:
                    continue
                paragraph = document.add_paragraph()
                run = paragraph.add_run(stripped)
                run.font.size = Pt(11)
                run.font.color.rgb = _NOTES_COLOR
        else:
            for _ in range(4):
                blank = document.add_paragraph()
                blank_run = blank.add_run("________________________________")
                blank_run.font.size = Pt(11)
                blank_run.font.color.rgb = _MUTED_COLOR

    def _add_footer_line(
        self,
        document: Document,
        slide: SlideSchema,
        deck: SlideDeckSchema,
    ) -> None:
        """Add a thin footer paragraph with deck title and slide number.

        Args:
            document: The DOCX document under construction.
            slide: The slide whose number is shown.
            deck: The deck whose title is shown.
        """
        footer = document.add_paragraph()
        footer_run = footer.add_run(
            f"{deck.title}  ·  Slide {slide.slide_number} / {deck.total}",
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = _MUTED_COLOR
        footer_run.italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    @staticmethod
    def _add_subtle_paragraph(document: Document, text: str) -> None:
        """Add a muted, italic paragraph used as a placeholder.

        Args:
            document: The DOCX document under construction.
            text: The paragraph content.
        """
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(11)
        run.italic = True
        run.font.color.rgb = _MUTED_COLOR

    @staticmethod
    def _add_page_break(document: Document) -> None:
        """Insert a hard page break before the next slide page.

        Args:
            document: The DOCX document under construction.
        """
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        break_element = run._element.makeelement(
            qn("w:br"),
            {qn("w:type"): "page"},
        )
        run._element.append(break_element)

    @staticmethod
    def _build_eyebrow(slide: SlideSchema) -> str:
        """Build the small label shown above the slide title.

        Args:
            slide: The slide schema.

        Returns:
            The label text, or an empty string when no label applies.
        """
        if slide.kind == SlideKind.LESSON_TITLE and slide.subtitle:
            return slide.subtitle.upper()
        if slide.kind == SlideKind.CONCLUSION:
            return "CONCLUSÃO"
        if slide.lesson_index is not None:
            return f"AULA {slide.lesson_index}"
        return ""
