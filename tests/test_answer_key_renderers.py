import io

import pytest
from docx import Document

from src.services import (
    AnswerKeyDocxService,
    AnswerKeyMarkdownService,
    AnswerKeyParserService,
)

SAMPLE_TEXT: str = (
    "Título: Gabarito de Lógica Digital\n"
    "Subtítulo: Disciplina | 2º Ano\n"
    "\n"
    "Parte 1: Portas Lógicas\n"
    "\n"
    "1. O que faz a porta NOT?\n"
    'R: Inverte a entrada (Y = <span class="overline">A</span>).\n'
    "\n"
    "2. Quando AND retorna 1?\n"
    "R: Apenas quando todas as entradas valem 1.\n"
)


@pytest.mark.asyncio
async def test_docx_renderer_produces_valid_office_file() -> None:
    """The DOCX renderer should emit a valid Office Open XML file."""
    parser: AnswerKeyParserService = AnswerKeyParserService()
    answer_key = await parser.parse(SAMPLE_TEXT)
    service: AnswerKeyDocxService = AnswerKeyDocxService()
    payload: bytes = await service.render(answer_key=answer_key)
    assert payload[:2] == b"PK"
    document = Document(io.BytesIO(payload))
    text: str = "\n".join(p.text for p in document.paragraphs)
    assert "Gabarito de Lógica Digital" in text
    assert "Parte 1: Portas Lógicas" in text


@pytest.mark.asyncio
async def test_docx_renderer_strips_inline_html() -> None:
    """Inline HTML in answers should not leak as raw markup into the DOCX."""
    parser: AnswerKeyParserService = AnswerKeyParserService()
    answer_key = await parser.parse(SAMPLE_TEXT)
    service: AnswerKeyDocxService = AnswerKeyDocxService()
    payload: bytes = await service.render(answer_key=answer_key)
    document = Document(io.BytesIO(payload))
    full_text: str = "\n".join(p.text for p in document.paragraphs)
    assert "<span" not in full_text
    assert "Inverte a entrada (Y = A)." in full_text


@pytest.mark.asyncio
async def test_markdown_renderer_uses_blockquote_for_answers() -> None:
    """Markdown output should format answers as `> Resposta:` blockquotes."""
    parser: AnswerKeyParserService = AnswerKeyParserService()
    answer_key = await parser.parse(SAMPLE_TEXT)
    service: AnswerKeyMarkdownService = AnswerKeyMarkdownService()
    output: str = await service.render(answer_key=answer_key)
    assert output.startswith("# Gabarito de Lógica Digital")
    assert "_Disciplina | 2º Ano_" in output
    assert "## Parte 1: Portas Lógicas" in output
    assert "> Resposta: Inverte a entrada (Y = A)." in output
    assert "<span" not in output
